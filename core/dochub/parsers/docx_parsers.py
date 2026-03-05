import base64
import os.path
import zipfile
from abc import ABC
from enum import StrEnum

import docx
from docx.document import Document as DocxDocument
from docx.drawing import Drawing
from docx.oxml.simpletypes import ST_Merge
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from dochub.parsers.base import BaseDocumentParser
from dochub.schemas import Chunk, Param, DataType, ChunkType
from utils import storage_utils
from dochub.utils.api_utils import gen_image_desc
from utils.i18n import I18NString, Language


class BaseDocxParser(BaseDocumentParser, ABC):
    target_content_type = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/wps-office.docx",
    ]
    target_file_ext = "docx"


class TableCellMergeModeEnum(StrEnum):
    KEEP_MERGED_CELLS = "Keep Merged Cells"
    SPLIT_MERGED_CELLS = "Split Merged Cells"


class GeneralDocxParser(BaseDocxParser):
    name = I18NString({
        Language.ZH: "通用",
        Language.EN: "General",
    })
    params = [
        Param(name="cell_merge_mode", display_name="表格合并单元格解析方式", required=True, data_type=DataType.STRING,
              default=TableCellMergeModeEnum.KEEP_MERGED_CELLS)
    ]

    def _parse_impl(self) -> None:
        doc = docx.Document(self.target.physical_path)

        # 将docx中的所有图片提取出来
        self._unzip_pictures()
        total_cnt = len(doc.paragraphs) + len(doc.tables)
        for part in self._progress_wrapper(doc.iter_inner_content(), total_cnt):
            if isinstance(part, Paragraph):
                # 解析段落
                self._parse_paragraph(part, doc)
            elif isinstance(part, Table):
                # 解析表格
                parse_mode = self.kwargs.get("cell_merge_mode", TableCellMergeModeEnum.KEEP_MERGED_CELLS)
                if parse_mode == TableCellMergeModeEnum.SPLIT_MERGED_CELLS:
                    self._split_table_merged_cells(part)
                else:
                    self._keep_table_merged_cells(part)

        self._report_progress(100)

    def _parse_paragraph(self, para: Paragraph, doc: DocxDocument) -> None:
        # 按顺序解析文本和图片
        text = None
        for elem in para.iter_inner_content():
            if isinstance(elem, Hyperlink) or not self.is_picture(elem):
                if text is None:
                    text = ""
                text += elem.text
            else:
                picture_info = self.get_picture_url_and_id(elem, doc)  # ("media/xxxx.png", "rIdx")
                if picture_info is None:
                    continue
                picture_url, picture_id = picture_info
                picture_name = picture_url.split("/")[1]
                out_file_path = f"{self._get_tmp_data_path()}/{picture_name}"
                with open(out_file_path, "rb") as f:
                    image_bytes = f.read()
                base64_str = base64.b64encode(image_bytes).decode("utf-8")
                img_mime = f"image/{picture_name.split(".")[-1]}"
                img_desc = gen_image_desc(base64_str, img_mime)
                chunk = Chunk(content=img_desc, type=ChunkType.IMAGE, metadata={
                    "image_base64": base64_str,
                    "image_mime": img_mime,
                })
                self._append_content_chunks(chunk)

        if text and len(text.strip()) > 0:
            self._append_content_chunks(Chunk(content=text, type=ChunkType.TEXT))

    def _unzip_pictures(self) -> None:
        """
        从`zip_file_path` 中将所有图片文件解压到 `out_dir` 路径下
        """
        with zipfile.ZipFile(self.target.physical_path) as zip_file:
            for f in zip_file.namelist():
                if f.startswith("word/media/") and not f.endswith("/"):
                    out_filename = f[f.rfind("/") + 1:]
                    out_file_path = os.path.join(self._get_tmp_data_path(), out_filename)
                    with zip_file.open(f) as file:
                        contents = file.read()
                        with open(out_file_path, "wb+") as out_file:
                            out_file.write(contents)

    def is_picture(self, run: Run) -> bool:
        for child in run.iter_inner_content():
            if isinstance(child, Drawing):
                return True
        return False

    def get_picture_url_and_id(self, run: Run, doc: DocxDocument):
        """
        需要假设一个 Run 里面只能有一张图片, 这个假设通常是成立的
        """
        for image_rid in run.element.xpath(".//a:blip/@r:embed"):
            return doc.part.target_ref(image_rid), image_rid

    def _split_table_merged_cells(self, table: Table) -> None:
        """
        一个 Chunk 代表一张表格
        会把合并的单元格拆开，内容复制多份。
        表格里面的图片无法解析出来，会直接忽略掉
        """
        content = ""
        for row in table.rows:
            text = "|"
            for cell in row.cells:
                text += cell.text
                text += "|"
            text += "\n"
            content += text

        if content:
            tbl_metadata = {
                "rows": len(table.rows),
                "columns": len(table.columns),
            }
            chunk = Chunk(content=content, type=ChunkType.TABLE, metadata=tbl_metadata)
            self._append_content_chunks(chunk)

    def _keep_table_merged_cells(self, table: Table) -> None:
        """
        合并的单元格被当成一个格子解析成一个grid，记录它的位置、大小和内容。
        未合并的单元格同样被解析成一个grid，grid内容同上
        表格里面的图片无法解析出来，会直接忽略掉
        """
        content = ""
        grids = []
        for tr in table._tbl.tr_lst:
            text = "|"
            for tc in tr.tc_lst:
                if tc.vMerge == ST_Merge.CONTINUE:  # 表明和上面一行的对应格子合并了
                    pass
                else:
                    try:
                        grid = {
                            "row_position": tc.top,
                            "col_position": tc.left,
                            "rows": tc.bottom - tc.top,
                            "cols": tc.right - tc.left
                        }
                    except:
                        grid = {}
                    grid["text"] = "".join([ice.text for ice in tc.inner_content_elements if hasattr(ice, "text") and ice.text])
                    grids.append(grid)
                    text += grid["text"]
                    text += "|"
            text += "\n"
            content += text

        if content:
            tbl_metadata = {
                "rows": len(table.rows),
                "columns": len(table.columns),
                "grids": grids,
            }
            chunk = Chunk(content=content, type=ChunkType.TABLE, metadata=tbl_metadata)
            self._append_content_chunks(chunk)
